import pandas as pd
import os
import random
from datetime import datetime
from docx import Document
from docx2pdf import convert
import inflect

p = inflect.engine()

def replace_placeholders(doc, replacements):
    # Replace in paragraphs (including runs)
    for pgraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in pgraph.text:
                for run in pgraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def generate_invoice(row, template_path, output_docx, output_pdf):
    doc = Document(template_path)

    # Parse date safely
    date_obj = pd.to_datetime(row["Date"], dayfirst=True, errors="coerce")
    if pd.isnull(date_obj):
        date_obj = datetime.today()

    # Generate invoice number
    invoice_number = date_obj.strftime("%Y%m%d") + str(random.randint(1000, 9999))

    # Convert amount to words
    amount_words = p.number_to_words(int(row["Amount"]), andword="") + " Only"
    amount_words = amount_words.title()

    # Prepare replacements
    replacements = {
        "<Buyer Name>": str(row["Buyer Name"]),
        "<Buyer Address>": str(row["Buyer Address"]),
        "<Buyer GSTIN>": str(row["Buyer GSTIN"]),
        "<Purpose>": str(row["Purpose"]),
        "<Amount>": str(row["Amount"]),
        "<Date>": date_obj.strftime("%d-%m-%Y"),
        "<Invoice Number>": invoice_number,
        "<Amount Words>": amount_words
    }

    # Replace placeholders everywhere
    replace_placeholders(doc, replacements)

    # Save Word invoice
    doc.save(output_docx)

    # Convert to PDF (Windows + Word required)
    try:
        convert(output_docx, output_pdf)
    except Exception as e:
        print(f"PDF conversion failed: {e}. Word file saved at {output_docx}")

def generate_invoices_from_excel(excel_file, template_path, output_dir="output"):
    df = pd.read_excel(excel_file)
    os.makedirs(output_dir, exist_ok=True)

    for idx, row in df.iterrows():
        buyer_name = str(row["Buyer Name"]).replace(" ", "_")
        docx_file = os.path.join(output_dir, f"Invoice_{buyer_name}_{idx+1}.docx")
        pdf_file = os.path.join(output_dir, f"Invoice_{buyer_name}_{idx+1}.pdf")
        generate_invoice(row, template_path, docx_file, pdf_file)
        print(f"Created: {pdf_file}")

if __name__ == "__main__":
    excel_file = r"C:\Users\debjy\PycharmProjects\HelloWorld\data\invoices.xlsx"
    template_file = r"C:\Users\debjy\PycharmProjects\HelloWorld\templates\InvoiceGSTDecl_Template.docx"
    generate_invoices_from_excel(excel_file, template_file)