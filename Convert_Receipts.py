from copy import deepcopy
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK


BASE = Path(r"C:\Users\debjy\OneDrive\Documents\GGLCA")

EXCEL_FILE = BASE / "Audit Checklist for FY 2025-26 - GGLCA.xlsx"
TEMPLATE_FILE = BASE / "CASH_RECEIPT Format.docx"
OUTPUT_FILE = BASE / "CASH_RECEIPTS_GENERATED.docx"


REQUIRED_COLUMNS = {
    "date": "Date",
    "amount": "Amount",
    "purpose": "Purpose",
    "particulars": "Particulars",
}


ONES = [
    "", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
    "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
    "Seventeen", "Eighteen", "Nineteen",
]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]


def two_digits_to_words(n: int) -> str:
    if n < 20:
        return ONES[n]
    return (TENS[n // 10] + " " + ONES[n % 10]).strip()


def three_digits_to_words(n: int) -> str:
    hundred = n // 100
    rest = n % 100
    parts = []
    if hundred:
        parts.append(ONES[hundred] + " Hundred")
    if rest:
        parts.append(two_digits_to_words(rest))
    return " ".join(parts)


def integer_to_indian_words(n: int) -> str:
    if n == 0:
        return "Zero"

    parts = []

    crore = n // 10000000
    n %= 10000000
    lakh = n // 100000
    n %= 100000
    thousand = n // 1000
    n %= 1000

    if crore:
        parts.append(three_digits_to_words(crore) + " Crore")
    if lakh:
        parts.append(three_digits_to_words(lakh) + " Lakh")
    if thousand:
        parts.append(three_digits_to_words(thousand) + " Thousand")
    if n:
        parts.append(three_digits_to_words(n))

    return " ".join(parts)


def amount_to_words(value) -> str:
    amount = Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rupees = int(amount)
    paise = int((amount - rupees) * 100)

    words = integer_to_indian_words(rupees) + " Rupees"
    if paise:
        words += " and " + integer_to_indian_words(paise) + " Paise"

    return words + " Only"


def format_date(value) -> str:
    if pd.isna(value):
        return ""

    if isinstance(value, pd.Timestamp):
        return value.strftime("%d-%m-%Y")

    if isinstance(value, (datetime, date)):
        return value.strftime("%d-%m-%Y")

    parsed = pd.to_datetime(value, errors="coerce")
    if pd.notna(parsed):
        return parsed.strftime("%d-%m-%Y")

    return str(value)


def clean_text(value) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def find_columns(df: pd.DataFrame) -> dict:
    normalized = {str(col).strip().lower(): col for col in df.columns}
    found = {}

    for key, display_name in REQUIRED_COLUMNS.items():
        lookup = display_name.lower()
        if lookup not in normalized:
            raise ValueError(f"Missing required column: {display_name}")
        found[key] = normalized[lookup]

    return found


def replace_text_in_paragraph(paragraph, replacements: dict):
    original_text = paragraph.text
    new_text = original_text

    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == original_text:
        return

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""

    first_run.text = new_text


def replace_text_everywhere(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def append_document_body(target: Document, source: Document):
    target_body = target.element.body
    source_body = source.element.body

    for element in source_body:
        if element.tag.endswith("}sectPr"):
            continue
        target_body.append(deepcopy(element))


def add_spacing_or_page_break(doc: Document, row_index: int):
    paragraph = doc.add_paragraph()

    if (row_index + 1) % 3 == 0:
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    else:
        paragraph.paragraph_format.space_after = 12


def main():
    df = pd.read_excel(EXCEL_FILE)
    columns = find_columns(df)

    rows = []
    for _, row in df.iterrows():
        amount = row[columns["amount"]]
        if pd.isna(amount):
            continue

        rows.append({
            "<Date>": format_date(row[columns["date"]]),
            "<Amount>": clean_text(amount),
            "<Amount in Words>": amount_to_words(amount),
            "<Purpose>": clean_text(row[columns["purpose"]]),
            "<Particulars>": clean_text(row[columns["particulars"]]),
        })

    if not rows:
        raise ValueError("No valid receipt rows found in the Excel file.")

    output_doc = Document(TEMPLATE_FILE)

    for element in list(output_doc.element.body):
        if not element.tag.endswith("}sectPr"):
            output_doc.element.body.remove(element)

    for index, replacements in enumerate(rows):
        block_doc = Document(TEMPLATE_FILE)
        replace_text_everywhere(block_doc, replacements)
        append_document_body(output_doc, block_doc)

        if index != len(rows) - 1:
            add_spacing_or_page_break(output_doc, index)

    output_doc.save(OUTPUT_FILE)
    print(f"Created: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()